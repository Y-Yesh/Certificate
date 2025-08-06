from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import os
from datetime import datetime

def debug_document_content(doc):
    """
    Debug function to see what's actually in the document
    """
    print("=== Document Content Debug ===")
    print(f"Number of paragraphs: {len(doc.paragraphs)}")
    print(f"Number of tables: {len(doc.tables)}")
    
    print("\n--- All Paragraphs ---")
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text:
            print(f"Paragraph {i}: '{text}'")
    
    print("\n--- All Tables ---")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}:")
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    print(f"  Cell {row_idx},{cell_idx}: '{text}'")

def set_paragraph_formatting(paragraph, alignment=None, font_size=None, font_name=None, bold=None, character_spacing=None):
    """
    Set formatting for a paragraph including alignment, font properties, and character spacing.
    """
    if alignment:
        if alignment.lower() == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment.lower() == 'left':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment.lower() == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment.lower() == 'justify':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Apply formatting to all runs in the paragraph
    for run in paragraph.runs:
        if font_size:
            run.font.size = Pt(font_size)
        if font_name:
            run.font.name = font_name
        if bold is not None:
            run.font.bold = bold
        if character_spacing:
            # Character spacing is set in points (1 point = 1/72 inch)
            run.font.spacing = Pt(character_spacing)

def replace_name_placeholder(doc, name, alignment='center', font_size=24, font_name='Arial', bold=True, character_spacing=0):
    """
    Replace the Name placeholder in the document with the provided name.
    """
    replacements_made = 0
    
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "Ally  Farah":
            paragraph.text = name
            set_paragraph_formatting(paragraph, alignment, font_size, font_name, bold, character_spacing)
            replacements_made += 1
            print(f"Replaced name in paragraph {i} with formatting: {alignment}, {font_size}pt, {font_name}, bold={bold}")
    
    return replacements_made

def replace_company_placeholder(doc, company, alignment='center', font_size=14, font_name='Arial', bold=False, character_spacing=0):
    """
    Replace the Company placeholder in the document with the provided company.
    """
    replacements_made = 0
    
    for i, paragraph in enumerate(doc.paragraphs):
        if "JEBSEN GROUP" in paragraph.text:
            paragraph.text = paragraph.text.replace("JEBSEN GROUP", company)
            set_paragraph_formatting(paragraph, alignment, font_size, font_name, bold, character_spacing)
            replacements_made += 1
            print(f"Replaced company in paragraph {i} with formatting: {alignment}, {font_size}pt, {font_name}, bold={bold}")
    
    return replacements_made

def replace_date_placeholder(doc, date, alignment='center', font_size=12, font_name='Arial', bold=False, character_spacing=0):
    """
    Replace the Date placeholder in the document with the provided date.
    """
    replacements_made = 0
    
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "AUGUST 7 – 8 , 2025":
            paragraph.text = date
            set_paragraph_formatting(paragraph, alignment, font_size, font_name, bold, character_spacing)
            replacements_made += 1
            print(f"Replaced date in paragraph {i} with formatting: {alignment}, {font_size}pt, {font_name}, bold={bold}")
    
    return replacements_made

def main():
    # Load the Word document
    doc_path = "t.docx"
    
    if not os.path.exists(doc_path):
        print(f"Error: {doc_path} not found!")
        return None
    
    # Load the document
    doc = Document(doc_path)
    
    # Debug: Show document content to find actual placeholders
    debug_document_content(doc)
    
    # ===== CONFIGURATION SECTION =====
    # Customize these values as needed
    
    # Certificate content
    name = "John Doe"
    company = "ACME Corporation"
    date = "DECEMBER 15 - 16, 2024"
    
    # Name formatting options
    name_format = {
        'alignment': 'center',      # 'left', 'center', 'right', 'justify'
        'font_size': 60,            # Font size in points
        'font_name': 'Arial',       # Font family
        'bold': True,               # True/False
        'character_spacing': 2      # Spacing in points (0 = normal)
    }
    
    # Company formatting options
    company_format = {
        'alignment': 'center',
        'font_size': 16,
        'font_name': 'Poppins',
        'bold': False,
        'character_spacing': 101
    }
    
    # Date formatting options
    date_format = {
        'alignment': 'center',
        'font_size': 14,
        'font_name': 'Garet',
        'bold': False,
        'character_spacing': 0
    }
    
    # Output filename
    output_filename = "certificate_formatted.docx"
    
    # ===== END CONFIGURATION =====
    
    print(f"\n=== Replacing Placeholders with Custom Formatting ===")
    
    # Name formatting: centered, large font, bold, with character spacing
    name_replacements = replace_name_placeholder(
        doc, name, 
        **name_format
    )
    
    # Company formatting: centered, medium font, not bold
    company_replacements = replace_company_placeholder(
        doc, company, 
        **company_format
    )
    
    # Date formatting: centered, smaller font, not bold
    date_replacements = replace_date_placeholder(
        doc, date, 
        **date_format
    )
    
    total_replacements = name_replacements + company_replacements + date_replacements
    
    if total_replacements > 0:
        # Save the modified document with a different name to avoid permission issues
        output_path = output_filename
        doc.save(output_path)
        print(f"\nCertificate with all placeholders filled and formatted saved as: {output_path}")
        print(f"Total replacements made: {total_replacements}")
    else:
        print("No placeholders found in the document.")
    
    return None

if __name__ == "__main__":
    main()
