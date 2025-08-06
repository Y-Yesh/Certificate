#!/usr/bin/env python3
"""
Font Demo Script - Shows how to check font availability and use fonts safely
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def get_available_fonts():
    """Get a list of commonly available fonts on Windows systems."""
    common_fonts = [
        'Arial', 'Arial Black', 'Arial Narrow', 'Arial Unicode MS',
        'Calibri', 'Cambria', 'Cambria Math', 'Candara',
        'Comic Sans MS', 'Consolas', 'Constantia', 'Corbel',
        'Courier New', 'Georgia', 'Impact', 'Lucida Console',
        'Lucida Sans Unicode', 'Microsoft Sans Serif', 'Palatino Linotype',
        'Segoe UI', 'Tahoma', 'Times New Roman', 'Trebuchet MS',
        'Verdana', 'Webdings', 'Wingdings',
        'Helvetica', 'Times', 'Courier', 'Symbol', 'ZapfDingbats',
        'Bookman Old Style', 'Century Gothic', 'Century Schoolbook',
        'Franklin Gothic Medium', 'Garamond', 'MS Gothic', 'MS Mincho',
        'MS PGothic', 'MS PMincho', 'MS Reference Sans Serif',
        'MS Reference Specialty', 'Rockwell', 'Stencil', 'Tw Cen MT'
    ]
    return common_fonts

def get_system_fonts():
    """Get actual system fonts using matplotlib if available."""
    try:
        import matplotlib.font_manager as fm
        font_list = [f.name for f in fm.fontManager.ttflist]
        unique_fonts = sorted(list(set(font_list)))
        return unique_fonts
    except ImportError:
        print("matplotlib not available. Using basic font list.")
        return get_available_fonts()

def check_font_availability(font_name):
    """Check if a font is in the basic list."""
    available_fonts = get_available_fonts()
    return font_name in available_fonts

def check_font_availability_advanced(font_name):
    """Check if a font is available using system enumeration."""
    system_fonts = get_system_fonts()
    return font_name in system_fonts

def safe_font_name(font_name, fallback='Arial'):
    """Safely set a font name with fallback."""
    if check_font_availability_advanced(font_name):
        return font_name
    else:
        print(f"Warning: Font '{font_name}' may not be available. Using fallback '{fallback}'.")
        return fallback

def create_font_demo_document():
    """Create a demo document showing different fonts."""
    doc = Document()
    
    # Title
    title = doc.add_heading('Font Demo Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Test different fonts
    test_fonts = [
        ('Arial', 'Arial is a clean, readable sans-serif font'),
        ('Times New Roman', 'Times New Roman is a classic serif font'),
        ('Calibri', 'Calibri is a modern sans-serif font'),
        ('Georgia', 'Georgia is an elegant serif font'),
        ('Comic Sans MS', 'Comic Sans MS is a casual, friendly font'),
        ('Courier New', 'Courier New is a monospace font'),
        ('Impact', 'Impact is a bold, attention-grabbing font'),
        ('Poppins', 'Poppins might not be available (will fallback to Arial)'),
        ('Garet', 'Garet might not be available (will fallback to Arial)')
    ]
    
    for font_name, description in test_fonts:
        # Check availability
        is_available = check_font_availability_advanced(font_name)
        status = "✓" if is_available else "✗"
        
        # Create paragraph
        p = doc.add_paragraph()
        p.add_run(f"{status} {font_name}: ").bold = True
        p.add_run(description)
        
        # Apply font to the description part
        safe_font = safe_font_name(font_name)
        for run in p.runs:
            if not run.bold:  # Only apply to the description part
                run.font.name = safe_font
                run.font.size = Pt(12)
    
    # Add a section about font checking
    doc.add_heading('Font Availability Information', level=1)
    
    info_para = doc.add_paragraph()
    info_para.add_run("Font checking methods:\n").bold = True
    info_para.add_run("1. Basic check: Uses a predefined list of common fonts\n")
    info_para.add_run("2. Advanced check: Uses matplotlib to enumerate actual system fonts\n")
    info_para.add_run("3. Safe font usage: Automatically falls back to Arial if font is unavailable\n")
    
    # Save the document
    output_file = "font_demo_output.docx"
    doc.save(output_file)
    print(f"Font demo document saved as: {output_file}")
    return output_file

def main():
    print("=== Font Availability and Usage Demo ===\n")
    
    # Test specific fonts
    test_fonts = ['Arial', 'Calibri', 'Times New Roman', 'Poppins', 'Garet', 'Comic Sans MS', 'NonExistentFont']
    
    print("Font Availability Check:")
    print("-" * 50)
    
    for font in test_fonts:
        basic_check = check_font_availability(font)
        advanced_check = check_font_availability_advanced(font)
        
        basic_status = "✓" if basic_check else "✗"
        advanced_status = "✓" if advanced_check else "✗"
        
        print(f"{font:20} | Basic: {basic_status} | Advanced: {advanced_status}")
    
    print(f"\nTotal fonts in basic list: {len(get_available_fonts())}")
    system_fonts = get_system_fonts()
    print(f"Total fonts on system: {len(system_fonts)}")
    
    # Show some system fonts
    print(f"\nFirst 15 system fonts:")
    for i, font in enumerate(system_fonts[:15]):
        print(f"  {i+1:2d}. {font}")
    
    # Create demo document
    print(f"\nCreating font demo document...")
    create_font_demo_document()
    
    print(f"\n=== Font Usage Tips ===")
    print("1. Always use safe_font_name() to avoid font errors")
    print("2. Common safe fonts: Arial, Calibri, Times New Roman, Georgia")
    print("3. Check font availability before using custom fonts")
    print("4. Provide fallback fonts for better compatibility")

if __name__ == "__main__":
    main() 